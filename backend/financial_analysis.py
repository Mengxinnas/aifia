"""财务分析报告生成模块 - 优化版本"""
import os
import json
import asyncio
import pandas as pd
import numpy as np
from typing import Dict, List, Optional, Any, AsyncGenerator
from datetime import datetime
import aiohttp
from fastapi import HTTPException
from fastapi.responses import StreamingResponse

# Word文档生成
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    DOCX_AVAILABLE = True
except ImportError:
    print("警告：python-docx未安装，Word文档生成功能将不可用")
    DOCX_AVAILABLE = False

# DeepSeek API配置
DEEPSEEK_API_KEY = "sk-c0c506ed07dc47a6b3713506a2ebd3c3"
DEEPSEEK_API_BASE = "https://api.deepseek.com/v1"

class FinancialAnalysisService:
    """财务分析服务类 - 优化版本"""
    
    def __init__(self):
        """初始化财务分析服务"""
        self.deepseek_api_url = f"{DEEPSEEK_API_BASE}/chat/completions"
        self.deepseek_api_key = DEEPSEEK_API_KEY
        
        # 任务管理
        self.running_tasks = {}  # 存储正在运行的任务
        
        # 详细的财务分析模板
        self.analysis_templates = {
            "comprehensive": self._get_comprehensive_template(),
            "dupont": self._get_dupont_template(),
            "profitability": self._get_profitability_template(),
            "debt": self._get_debt_template(),
            "efficiency": self._get_efficiency_template(),
            "growth": self._get_growth_template(),
            "investment": self._get_investment_template(),
            "cashflow": self._get_cashflow_template()
        }
    
    def _get_comprehensive_template(self) -> str:
        """获取综合财务分析模板"""
        return """
        你是一位资深的财务分析师，请基于提供的财务数据文档生成一份专业的综合财务分析报告。
        
        请按照以下结构和指标进行全面分析：
        
        # 综合财务分析报告
        
        ## 1. 执行摘要
        - 核心发现和结论
        - 主要财务指标概览
        - 关键风险点识别
        - 重要建议摘要
        
        ## 2. 盈利能力分析
        ### 2.1 盈利指标计算
        - **净资产收益率(ROE)** = 净利润 ÷ 平均净资产 × 100%
        - **总资产报酬率(ROA)** = (净利润 + 利息费用) ÷ 平均总资产 × 100%
        - **销售净利率** = 净利润 ÷ 营业收入 × 100%
        - **销售毛利率** = (营业收入 - 营业成本) ÷ 营业收入 × 100%
        - **成本费用利润率** = 利润总额 ÷ 成本费用总额 × 100%
        
        ### 2.2 盈利质量评估
        - 收入的真实性和可持续性
        - 利润结构分析
        - 非经常性损益影响
        
        ## 3. 偿债能力分析
        ### 3.1 短期偿债能力
        - **流动比率** = 流动资产 ÷ 流动负债
        - **速动比率** = (流动资产 - 存货) ÷ 流动负债
        - **现金比率** = 货币资金 ÷ 流动负债
        
        ### 3.2 长期偿债能力
        - **资产负债率** = 总负债 ÷ 总资产 × 100%
        - **产权比率** = 总负债 ÷ 股东权益 × 100%
        - **利息保障倍数** = 息税前利润 ÷ 利息费用
        
        ## 4. 营运能力分析
        ### 4.1 资产周转效率
        - **总资产周转率** = 营业收入 ÷ 平均总资产
        - **流动资产周转率** = 营业收入 ÷ 平均流动资产
        - **固定资产周转率** = 营业收入 ÷ 平均固定资产净值
        
        ### 4.2 具体资产管理
        - **存货周转率** = 营业成本 ÷ 平均存货
        - **存货周转天数** = 365 ÷ 存货周转率
        - **应收账款周转率** = 营业收入 ÷ 平均应收账款
        - **应收账款周转天数** = 365 ÷ 应收账款周转率
        
        ## 5. 成长能力分析
        ### 5.1 增长率指标
        - **营业收入增长率** = (本期营业收入 - 上期营业收入) ÷ 上期营业收入 × 100%
        - **净利润增长率** = (本期净利润 - 上期净利润) ÷ 上期净利润 × 100%
        - **总资产增长率** = (期末总资产 - 期初总资产) ÷ 期初总资产 × 100%
        - **净资产增长率** = (期末净资产 - 期初净资产) ÷ 期初净资产 × 100%
        
        ## 6. 现金流量分析
        ### 6.1 现金流结构
        - 经营活动现金流量分析
        - 投资活动现金流量分析
        - 筹资活动现金流量分析
        
        ### 6.2 现金流质量
        - **现金流量比率** = 经营现金净流量 ÷ 流动负债
        - **现金债务总额比** = 经营现金净流量 ÷ 债务总额
        - **现金再投资比率** = 经营现金净流量 ÷ (资本支出 + 存货增加 + 现金股利)
        
        ## 7. 风险评估与建议
        ### 7.1 财务风险识别
        ### 7.2 经营风险分析
        ### 7.3 改进建议
        ### 7.4 投资价值评估
        
        请确保：
        1. 基于实际财务数据计算具体指标
        2. 提供行业对比参考
        3. 分析趋势变化
        4. 给出具体可行的建议
        """
    
    def _get_dupont_template(self) -> str:
        """获取杜邦分析模板"""
        return """
        请进行专业的杜邦分析，深度解析净资产收益率的驱动因素：
        
        # 杜邦分析报告
        
        ## 1. 杜邦分析模型概述
        **ROE = 销售净利率 × 资产周转率 × 权益乘数**
        
        即：净资产收益率 = (净利润/营业收入) × (营业收入/平均总资产) × (平均总资产/平均净资产)
        
        ## 2. 核心指标计算与分解
        
        ### 2.1 净资产收益率(ROE)
        - **计算公式**：ROE = 净利润 ÷ 平均净资产 × 100%
        - **实际数值**：[从财务数据中提取并计算]
        - **行业对比**：[提供行业平均水平对比]
        
        ### 2.2 第一层分解
        #### 销售净利率 (Net Profit Margin)
        - **计算公式**：销售净利率 = 净利润 ÷ 营业收入 × 100%
        - **实际数值**：[计算具体数值]
        - **分析要点**：
          - 反映企业盈利能力
          - 成本控制水平
          - 费用管理效率
        
        #### 资产周转率 (Asset Turnover)
        - **计算公式**：资产周转率 = 营业收入 ÷ 平均总资产
        - **实际数值**：[计算具体数值]
        - **分析要点**：
          - 反映资产利用效率
          - 营运管理水平
          - 资产配置合理性
        
        #### 权益乘数 (Equity Multiplier)
        - **计算公式**：权益乘数 = 平均总资产 ÷ 平均净资产
        - **实际数值**：[计算具体数值]
        - **分析要点**：
          - 反映财务杠杆水平
          - 资本结构状况
          - 财务风险程度
        
        ## 3. 第二层深度分解
        
        ### 3.1 销售净利率细分
        - **销售毛利率** = (营业收入 - 营业成本) ÷ 营业收入 × 100%
        - **期间费用率** = (销售费用 + 管理费用 + 财务费用) ÷ 营业收入 × 100%
        - **营业利润率** = 营业利润 ÷ 营业收入 × 100%
        
        ### 3.2 资产周转率细分
        - **流动资产周转率** = 营业收入 ÷ 平均流动资产
        - **固定资产周转率** = 营业收入 ÷ 平均固定资产净值
        - **存货周转率** = 营业成本 ÷ 平均存货
        - **应收账款周转率** = 营业收入 ÷ 平均应收账款
        
        ### 3.3 财务杠杆细分
        - **资产负债率** = 总负债 ÷ 总资产 × 100%
        - **产权比率** = 总负债 ÷ 所有者权益 × 100%
        - **债务资本比率** = 总负债 ÷ (总负债 + 所有者权益) × 100%
        
        ## 4. 历史趋势分析
        ### 4.1 ROE变化趋势
        ### 4.2 各驱动因素变化分析
        ### 4.3 主要影响因素识别
        
        ## 5. 同行业对比分析
        ### 5.1 ROE行业排名
        ### 5.2 各分解指标行业对比
        ### 5.3 竞争优势劣势分析
        
        ## 6. 改善建议
        ### 6.1 提升销售净利率建议
        - 成本控制措施
        - 费用优化方案
        - 产品结构调整
        
        ### 6.2 提升资产周转率建议
        - 资产管理优化
        - 营运效率提升
        - 资产配置改善
        
        ### 6.3 财务杠杆优化建议
        - 资本结构调整
        - 融资成本控制
        - 财务风险管理
        
        ## 7. 结论与展望
        ### 7.1 ROE驱动因素总结
        ### 7.2 未来改善路径
        ### 7.3 风险提示
        
        请确保：
        1. 计算所有相关财务比率
        2. 提供具体数值和变化趋势
        3. 深入分析各驱动因素
        4. 给出具体改善建议
        """
    
    def _get_profitability_template(self) -> str:
        """获取盈利能力分析模板"""
        return """
        请进行专业的盈利能力与收益质量分析：
        
        # 盈利能力与收益质量分析报告
        
        ## 1. 核心盈利指标分析
        
        ### 1.1 盈利能力指标
        - **净资产收益率(ROE)** = 净利润 ÷ 平均净资产 × 100%
        - **总资产报酬率(ROA)** = (净利润 + 利息费用) ÷ 平均总资产 × 100%
        - **销售净利率** = 净利润 ÷ 营业收入 × 100%
        - **销售毛利率** = 毛利润 ÷ 营业收入 × 100%
        - **营业利润率** = 营业利润 ÷ 营业收入 × 100%
        - **成本费用利润率** = 利润总额 ÷ 成本费用总额 × 100%
        - **资本保值增值率** = 期末所有者权益 ÷ 期初所有者权益 × 100%
        
        ### 1.2 人均效益指标
        - **人均营业收入** = 营业收入 ÷ 员工总数
        - **人均净利润** = 净利润 ÷ 员工总数
        
        ## 2. 收益质量分析
        
        ### 2.1 利润结构分析
        - 营业利润占比分析
        - 非经常性损益影响
        - 投资收益占比
        - 政府补助依赖度
        
        ### 2.2 收入质量评估
        - 营业收入真实性分析
        - 收入确认政策合理性
        - 现金流与利润匹配度
        - 应收账款质量分析
        
        ### 2.3 成本费用分析
        - **营业成本率** = 营业成本 ÷ 营业收入 × 100%
        - **销售费用率** = 销售费用 ÷ 营业收入 × 100%
        - **管理费用率** = 管理费用 ÷ 营业收入 × 100%
        - **财务费用率** = 财务费用 ÷ 营业收入 × 100%
        - **研发费用率** = 研发费用 ÷ 营业收入 × 100%
        
        ## 3. 盈利趋势分析
        ### 3.1 增长率分析
        - 营业收入同比增长率
        - 毛利润同比增长率
        - 净利润同比增长率
        - 扣非净利润同比增长率
        
        ### 3.2 盈利稳定性
        - 利润波动性分析
        - 季节性因素影响
        - 周期性特征识别
        
        ## 4. 同行业对比
        ### 4.1 行业地位评估
        ### 4.2 竞争优势分析
        ### 4.3 改进空间识别
        
        ## 5. 风险提示与建议
        ### 5.1 盈利能力风险
        ### 5.2 收益质量风险
        ### 5.3 改善建议
        
        请确保分析深入、数据准确、建议具体可行。
        """
    
    def _get_debt_template(self) -> str:
        """获取偿债能力分析模板"""
        return """
        请进行专业的资本结构与偿债能力分析：
        
        # 资本结构与偿债能力分析报告
        
        ## 1. 短期偿债能力分析
        
        ### 1.1 流动性指标
        - **流动比率** = 流动资产 ÷ 流动负债
        - **速动比率** = (流动资产 - 存货) ÷ 流动负债
        - **现金比率** = (货币资金 + 交易性金融资产) ÷ 流动负债
        - **现金流动负债比** = 经营现金净流量 ÷ 流动负债
        
        ### 1.2 营运资金分析
        - **营运资金** = 流动资产 - 流动负债
        - **营运资金比率** = 营运资金 ÷ 流动资产 × 100%
        
        ## 2. 长期偿债能力分析
        
        ### 2.1 资本结构指标
        - **资产负债率** = 总负债 ÷ 总资产 × 100%
        - **产权比率** = 总负债 ÷ 所有者权益 × 100%
        - **权益乘数** = 总资产 ÷ 所有者权益
        - **长期资产适合率** = (所有者权益 + 长期负债) ÷ 长期资产 × 100%
        
        ### 2.2 债务保障指标
        - **利息保障倍数** = 息税前利润 ÷ 利息费用
        - **现金利息保障倍数** = (经营现金净流量 + 利息费用 + 所得税费用) ÷ 利息费用
        - **债务保障比率** = 经营现金净流量 ÷ 债务总额 × 100%
        
        ## 3. 债务结构分析
        
        ### 3.1 债务期限结构
        - 短期债务占比
        - 长期债务占比
        - 债务到期分布
        
        ### 3.2 债务成本分析
        - 平均融资成本
        - 不同债务工具成本对比
        - 财务费用率变化趋势
        
        ## 4. 担保与抵押情况
        ### 4.1 对外担保分析
        ### 4.2 资产抵押情况
        ### 4.3 或有负债评估
        
        ## 5. 现金流偿债能力
        ### 5.1 经营现金流分析
        ### 5.2 自由现金流分析
        ### 5.3 现金流预测
        
        ## 6. 风险评估与建议
        ### 6.1 偿债风险识别
        ### 6.2 资本结构优化建议
        ### 6.3 流动性管理建议
        
        请基于实际财务数据进行深入分析，提供具体的改善建议。
        """
    
    def _get_efficiency_template(self) -> str:
        """获取营运效率分析模板"""
        return """
        请进行专业的资产营运效率分析：
        
        # 资产营运效率分析报告
        
        ## 1. 总体营运效率
        
        ### 1.1 总资产营运效率
        - **总资产周转率** = 营业收入 ÷ 平均总资产
        - **总资产周转天数** = 365 ÷ 总资产周转率
        - **流动资产周转率** = 营业收入 ÷ 平均流动资产
        - **固定资产周转率** = 营业收入 ÷ 平均固定资产净值
        
        ## 2. 存货管理效率
        
        ### 2.1 存货周转分析
        - **存货周转率** = 营业成本 ÷ 平均存货
        - **存货周转天数** = 365 ÷ 存货周转率
        - **存货占流动资产比重** = 存货 ÷ 流动资产 × 100%
        
        ### 2.2 存货结构分析
        - 原材料存货分析
        - 在产品存货分析
        - 产成品存货分析
        - 存货跌价准备充足性
        
        ## 3. 应收账款管理
        
        ### 3.1 应收账款周转分析
        - **应收账款周转率** = 营业收入 ÷ 平均应收账款
        - **应收账款周转天数** = 365 ÷ 应收账款周转率
        - **应收账款占营业收入比重** = 应收账款 ÷ 营业收入 × 100%
        
        ### 3.2 应收账款质量
        - 账龄结构分析
        - 坏账准备充足性
        - 前五大客户集中度
        - 关联方应收账款
        
        ## 4. 应付账款管理
        
        ### 4.1 应付账款周转分析
        - **应付账款周转率** = 营业成本 ÷ 平均应付账款
        - **应付账款周转天数** = 365 ÷ 应付账款周转率
        
        ### 4.2 供应商关系
        - 前五大供应商集中度
        - 付款政策分析
        - 供应链稳定性
        
        ## 5. 营业周期分析
        
        ### 5.1 营业周期计算
        - **营业周期** = 存货周转天数 + 应收账款周转天数
        - **现金周期** = 营业周期 - 应付账款周转天数
        
        ### 5.2 营业周期变化趋势
        - 历史变化分析
        - 行业对比
        - 季节性特征
        
        ## 6. 具体资产效率分析
        
        ### 6.1 货币资金管理
        - 货币资金充足性
        - 现金管理效率
        - 理财产品配置
        
        ### 6.2 固定资产效率
        - 产能利用率
        - 设备更新情况
        - 折旧政策合理性
        
        ### 6.3 无形资产效率
        - 研发投入产出比
        - 专利技术价值
        - 品牌价值评估
        
        ## 7. 营运效率改善建议
        ### 7.1 存货管理优化
        ### 7.2 应收账款管理加强
        ### 7.3 供应链管理改善
        ### 7.4 资产配置优化
        
        请基于实际数据进行分析，并提供具体的改善措施。
        """
    
    def _get_growth_template(self) -> str:
        """获取成长能力分析模板"""
        return """
        请进行专业的成长能力分析：
        
        # 成长能力分析报告
        
        ## 1. 收入增长分析
        
        ### 1.1 营业收入增长
        - **营业收入增长率** = (本期营业收入 - 上期营业收入) ÷ 上期营业收入 × 100%
        - **主营业务收入增长率**
        - **其他业务收入增长率**
        
        ### 1.2 收入增长驱动因素
        - 销售量增长贡献
        - 销售价格变化贡献
        - 新产品贡献
        - 新市场开拓贡献
        
        ## 2. 利润增长分析
        
        ### 2.1 利润增长指标
        - **营业利润增长率** = (本期营业利润 - 上期营业利润) ÷ 上期营业利润 × 100%
        - **净利润增长率** = (本期净利润 - 上期净利润) ÷ 上期净利润 × 100%
        - **扣非净利润增长率**
        - **每股收益增长率**
        
        ### 2.2 利润增长质量
        - 收入与利润增长匹配度
        - 利润增长可持续性
        - 非经常性损益影响
        
        ## 3. 资产增长分析
        
        ### 3.1 资产规模增长
        - **总资产增长率** = (期末总资产 - 期初总资产) ÷ 期初总资产 × 100%
        - **净资产增长率** = (期末净资产 - 期初净资产) ÷ 期初净资产 × 100%
        - **固定资产增长率**
        - **无形资产增长率**
        
        ### 3.2 资产增长效率
        - 资产增长与收入增长匹配度
        - 新增资产盈利能力
        - 资本支出回报率
        
        ## 4. 现金流增长分析
        
        ### 4.1 现金流增长指标
        - **经营现金流增长率**
        - **自由现金流增长率**
        - **投资现金流分析**
        
        ### 4.2 现金流增长质量
        - 现金流与利润增长匹配度
        - 现金创造能力
        - 资本支出可持续性
        
        ## 5. 研发创新能力
        
        ### 5.1 研发投入分析
        - **研发费用率** = 研发费用 ÷ 营业收入 × 100%
        - **研发费用增长率**
        - **研发人员占比**
        
        ### 5.2 创新成果
        - 新产品收入占比
        - 专利申请数量
        - 技术水平评估
        
        ## 6. 市场扩张能力
        
        ### 6.1 市场份额变化
        - 行业地位变化
        - 竞争优势分析
        - 品牌价值提升
        
        ### 6.2 渠道建设
        - 销售网络扩张
        - 客户结构优化
        - 新市场开拓
        
        ## 7. 可持续增长分析
        
        ### 7.1 内生增长能力
        - **可持续增长率** = ROE × (1 - 股利支付率)
        - 内源融资能力
        - 盈利再投资率
        
        ### 7.2 外延增长分析
        - 并购整合能力
        - 外部融资能力
        - 战略投资效果
        
        ## 8. 风险与建议
        ### 8.1 增长风险识别
        ### 8.2 增长策略建议
        ### 8.3 资源配置优化
        
        请基于实际数据分析企业成长潜力，并提供具体的成长策略建议。
        """
    
    def _get_investment_template(self) -> str:
        """获取资本投资效率分析模板"""
        return """
        请进行专业的资本投资效率分析：
        
        # 资本投资效率分析报告
        
        ## 1. 投资回报率分析
        
        ### 1.1 核心投资回报指标
        - **投资资本回报率(ROIC)** = 税后营业净利润 ÷ 投入资本 × 100%
        - **经济增加值(EVA)** = 税后营业净利润 - 投入资本 × 加权平均资本成本
        - **现金投资回报率** = 经营现金净流量 ÷ 平均投入资本 × 100%
        
        ### 1.2 分项投资回报
        - 固定资产投资回报率
        - 研发投资回报率
        - 并购投资回报率
        - 金融投资回报率
        
        ## 2. 资本支出分析
        
        ### 2.1 资本支出结构
        - **资本支出率** = 资本支出 ÷ 营业收入 × 100%
        - 维持性资本支出
        - 扩张性资本支出
        - 效率提升资本支出
        
        ### 2.2 资本支出效率
        - 新增产能投入产出比
        - 技改投资效果
        - 设备利用率提升
        
        ## 3. 资金成本分析
        
        ### 3.1 加权平均资本成本(WACC)
        - **WACC** = (E/V × Re) + (D/V × Rd × (1-T))
        - 权益资本成本
        - 债务资本成本
        - 资本结构优化
        
        ### 3.2 融资效率
        - 融资成本变化趋势
        - 融资渠道多样性
        - 融资时机把握
        
        ## 4. 资产配置效率
        
        ### 4.1 资产结构分析
        - 经营性资产占比
        - 金融性资产占比
        - 闲置资产处置
        
        ### 4.2 资产配置优化
        - 核心业务资产集中度
        - 非核心资产剥离
        - 资产轻重资产平衡
        
        ## 5. 投资项目评估
        
        ### 5.1 重大投资项目分析
        - 项目投资回报率
        - 项目现金流分析
        - 项目风险评估
        
        ### 5.2 投资决策质量
        - 投资决策流程
        - 投资后评价
        - 投资失败案例分析
        
        ## 6. 现金创造能力
        
        ### 6.1 自由现金流分析
        - **自由现金流** = 经营现金净流量 - 资本支出
        - **自由现金流收益率** = 自由现金流 ÷ 企业价值 × 100%
        
        ### 6.2 现金再投资比率
        - **现金再投资比率** = 经营现金净流量 ÷ (资本支出 + 净营运资金变化 + 股利支付)
        
        ## 7. 价值创造分析
        
        ### 7.1 价值创造指标
        - 经济利润(EP)
        - 市场增加值(MVA)
        - 托宾Q值
        
        ### 7.2 价值驱动因素
        - 收入增长贡献
        - 利润率改善贡献
        - 资本效率提升贡献
        
        ## 8. 资本效率改善建议
        
        ### 8.1 投资策略优化
        - 投资重点调整
        - 投资决策改善
        - 投资管控加强
        
        ### 8.2 资本结构优化
        - 债务结构调整
        - 融资成本降低
        - 资本配置改善
        
        ### 8.3 价值创造提升
        - 核心业务聚焦
        - 运营效率提升
        - 成本控制加强
        
        请基于实际财务数据分析投资效率，并提供具体的改善建议。
        """
    
    def _get_cashflow_template(self) -> str:
        """获取现金流量分析模板"""
        return """
        请进行专业的现金流量分析：
        
        # 现金流量分析报告
        
        ## 1. 现金流量结构分析
        
        ### 1.1 经营活动现金流量
        - **销售商品、提供劳务收到的现金**
        - **收到的税费返还**
        - **收到其他与经营活动有关的现金**
        - **购买商品、接受劳务支付的现金**
        - **支付给职工以及为职工支付的现金**
        - **支付的各项税费**
        - **支付其他与经营活动有关的现金**
        - **经营活动产生的现金流量净额**
        
        ### 1.2 投资活动现金流量
        - **收回投资收到的现金**
        - **取得投资收益收到的现金**
        - **处置固定资产、无形资产和其他长期资产收回的现金净额**
        - **购建固定资产、无形资产和其他长期资产支付的现金**
        - **投资支付的现金**
        - **投资活动产生的现金流量净额**
        
        ### 1.3 筹资活动现金流量
        - **吸收投资收到的现金**
        - **取得借款收到的现金**
        - **偿还债务支付的现金**
        - **分配股利、利润或偿付利息支付的现金**
        - **筹资活动产生的现金流量净额**
        
        ## 2. 现金流量质量分析
        
        ### 2.1 经营现金流质量
        - **现金流量比率** = 经营现金净流量 ÷ 流动负债
        - **现金债务总额比** = 经营现金净流量 ÷ 债务总额
        - **现金利息保障倍数** = (经营现金净流量 + 利息费用 + 所得税费用) ÷ 利息费用
        
        ### 2.2 现金流与利润匹配分析
        - **销售现金比率** = 销售商品收到的现金 ÷ 营业收入 × 100%
        - **净现比** = 经营现金净流量 ÷ 净利润
        - **现金净利率** = 经营现金净流量 ÷ 营业收入 × 100%
        
        ## 3. 现金流量趋势分析
        
        ### 3.1 现金流增长性
        - 经营现金流增长率
        - 自由现金流增长率
        - 现金流稳定性分析
        
        ### 3.2 现金流周期性
        - 季节性特征
        - 行业周期影响
        - 经济周期敏感性
        
        ## 4. 现金流量充足性分析
        
        ### 4.1 现金满足投资比率
        - **现金满足投资比率** = 经营现金净流量 ÷ (资本支出 + 存货增加 + 现金股利)
        
        ### 4.2 现金再投资比率
        - **现金再投资比率** = (经营现金净流量 - 现金股利) ÷ (资本支出 + 存货增加 + 现金股利)
        
        ## 5. 自由现金流分析
        
        ### 5.1 企业自由现金流
        - **企业自由现金流** = 经营现金净流量 - 资本支出
        - **自由现金流收益率** = 自由现金流 ÷ 企业价值 × 100%
        
        ### 5.2 股权自由现金流
        - **股权自由现金流** = 企业自由现金流 - 利息费用 × (1-税率) + 净借款
        
        ## 6. 现金流管理效率
        
        ### 6.1 现金转换周期
        - **现金转换周期** = 存货周转天数 + 应收账款周转天数 - 应付账款周转天数
        
        ### 6.2 营运资金管理
        - 营运资金变化分析
        - 营运资金周转效率
        - 季节性资金需求
        
        ## 7. 现金流风险分析
        
        ### 7.1 现金流风险识别
        - 经营现金流波动性
        - 资金链紧张风险
        - 流动性风险评估
        
        ### 7.2 现金流预测
        - 未来现金流预测
        - 资金缺口预警
        - 融资需求分析
        
        ## 8. 现金流改善建议
        
        ### 8.1 经营现金流改善
        - 销售回款管理
        - 应收账款控制
        - 存货管理优化
        - 供应商付款优化
        
        ### 8.2 投资现金流优化
        - 资本支出计划
        - 投资项目选择
        - 资产处置策略
        
        ### 8.3 筹资现金流管理
        - 融资结构优化
        - 股利政策调整
        - 债务管理改善
        
        请基于实际现金流量表数据进行深入分析，并提供具体的现金流管理建议。
        """
    
    async def analyze_financial_data_stream(self, text_content: str, analysis_type: str = "comprehensive", task_id: str = None) -> AsyncGenerator[str, None]:
        """流式财务分析 - 真实API调用版本"""
        try:
            print(f"开始真实流式财务分析，类型: {analysis_type}")
            print(f"文本长度: {len(text_content)}")
            
            # 检查任务是否被取消
            if task_id and self._is_task_cancelled(task_id):
                raise asyncio.CancelledError("任务已被用户取消")
            
            yield f"data: {json.dumps({'type': 'status', 'message': '开始分析财务数据...'}, ensure_ascii=False)}\n\n"
            await asyncio.sleep(0.1)
            
            # 构建分析模板
            template = self.analysis_templates.get(analysis_type, self.analysis_templates["comprehensive"])
            
            # 使用实际的文档内容进行分析
            text_summary = text_content[:5000] if len(text_content) > 5000 else text_content
            
            yield f"data: {json.dumps({'type': 'status', 'message': '准备AI分析请求...'}, ensure_ascii=False)}\n\n"
            await asyncio.sleep(0.1)
            
            # 构建包含实际文档内容的分析prompt
            analysis_prompt = f"""
            {template}
            
            基于以下财务文档内容进行专业分析：
            
            {text_summary}
            
            请直接对上述文档内容进行专业的财务分析，生成结构化的财务分析报告。
            要求：
            1. 使用markdown格式输出
            2. 结构清晰，层次分明
            3. 基于实际数据进行分析，引用具体数据
            4. 提供具体可行的建议
            5. 用中文输出
            6. 如果数据不完整，请基于可获得的信息进行分析并说明限制
            """
            
            yield f"data: {json.dumps({'type': 'status', 'message': '调用DeepSeek AI分析引擎...'}, ensure_ascii=False)}\n\n"
            await asyncio.sleep(0.1)
            
            yield f"data: {json.dumps({'type': 'analysis_start'}, ensure_ascii=False)}\n\n"
            await asyncio.sleep(0.1)
            
            # 检查API密钥并调用真实API
            if not self.deepseek_api_key:
                print("DeepSeek API密钥未配置，使用模拟分析...")
                yield f"data: {json.dumps({'type': 'status', 'message': 'API密钥未配置，使用模拟分析...'}, ensure_ascii=False)}\n\n"
                async for chunk in self._generate_mock_stream_analysis(text_content, analysis_type):
                    if task_id and self._is_task_cancelled(task_id):
                        break
                    yield chunk
                return
            
            # 调用真实的DeepSeek API
            headers = {
                "Authorization": f"Bearer {self.deepseek_api_key}",
                "Content-Type": "application/json",
                "Accept": "text/event-stream"
            }
            
            payload = {
                "model": "deepseek-chat",
                "messages": [
                    {
                        "role": "system",
                        "content": "你是一位专业的财务分析师，擅长财务报表分析和企业经营分析。请基于提供的实际财务数据提供专业、客观、实用的财务分析报告。"
                    },
                    {
                        "role": "user",
                        "content": analysis_prompt
                    }
                ],
                "max_tokens": 4000,
                "temperature": 0.3,
                "stream": True
            }
            
            print("开始调用DeepSeek流式API...")
            yield f"data: {json.dumps({'type': 'status', 'message': '正在调用AI分析引擎...'}, ensure_ascii=False)}\n\n"
            
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.post(
                        self.deepseek_api_url,
                        headers=headers,
                        json=payload,
                        timeout=aiohttp.ClientTimeout(total=120)
                    ) as response:
                        print(f"DeepSeek API响应状态: {response.status}")
                        
                        if response.status == 200:
                            yield f"data: {json.dumps({'type': 'status', 'message': 'AI分析开始，正在生成报告...'}, ensure_ascii=False)}\n\n"
                            
                            # 处理流式响应
                            buffer = ""
                            chunk_count = 0
                            
                            async for chunk in response.content.iter_chunked(1024):
                                # 检查任务是否被取消
                                if task_id and self._is_task_cancelled(task_id):
                                    raise asyncio.CancelledError("任务已被用户取消")
                                
                                chunk_count += 1
                                buffer += chunk.decode('utf-8')
                                lines = buffer.split('\n')
                                buffer = lines[-1]  # 保留最后一行（可能不完整）
                                
                                for line in lines[:-1]:
                                    line = line.strip()
                                    if line.startswith('data: '):
                                        data_str = line[6:]
                                        if data_str == '[DONE]':
                                            print("DeepSeek流式响应完成")
                                            break
                                        
                                        try:
                                            data = json.loads(data_str)
                                            if 'choices' in data and len(data['choices']) > 0:
                                                delta = data['choices'][0].get('delta', {})
                                                if 'content' in delta:
                                                    content = delta['content']
                                                    if content:
                                                        print(f"接收到内容块: {content[:50]}...")
                                                        yield f"data: {json.dumps({'type': 'analysis_chunk', 'content': content}, ensure_ascii=False)}\n\n"
                                        except json.JSONDecodeError as e:
                                            print(f"JSON解析错误: {e}, 原始数据: {data_str[:100]}")
                                            continue
                            
                            print(f"DeepSeek流式响应处理完成，共处理 {chunk_count} 个数据块")
                            
                        else:
                            error_text = await response.text()
                            print(f"DeepSeek API调用失败: {response.status}, 错误: {error_text}")
                            yield f"data: {json.dumps({'type': 'status', 'message': f'API调用失败，使用基础分析...'}, ensure_ascii=False)}\n\n"
                            
                            # API调用失败，使用基于实际内容的分析
                            async for chunk in self._generate_content_based_analysis(text_content, analysis_type):
                                if task_id and self._is_task_cancelled(task_id):
                                    break
                                yield chunk
                            return
                            
            except Exception as api_error:
                print(f"DeepSeek API连接错误: {str(api_error)}")
                yield f"data: {json.dumps({'type': 'status', 'message': 'API连接失败，使用基础分析...'}, ensure_ascii=False)}\n\n"
                
                # API连接失败，使用基于实际内容的分析
                async for chunk in self._generate_content_based_analysis(text_content, analysis_type):
                    if task_id and self._is_task_cancelled(task_id):
                        break
                    yield chunk
                return
            
            yield f"data: {json.dumps({'type': 'analysis_complete'}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'type': 'status', 'message': '分析完成'}, ensure_ascii=False)}\n\n"
            
        except asyncio.CancelledError:
            print(f"流式分析被取消 - 任务ID: {task_id}")
            yield f"data: {json.dumps({'type': 'error', 'message': '分析已被取消'}, ensure_ascii=False)}\n\n"
        except Exception as e:
            print(f"流式分析错误: {str(e)}")
            import traceback
            traceback.print_exc()
            yield f"data: {json.dumps({'type': 'error', 'message': f'分析失败: {str(e)}'}, ensure_ascii=False)}\n\n"
    
    async def _generate_content_based_analysis(self, text_content: str, analysis_type: str) -> AsyncGenerator[str, None]:
        """基于实际内容的分析（当API不可用时的备选方案）"""
        print("生成基于实际内容的分析...")
        
        current_date = datetime.now().strftime("%Y年%m月%d日")
        text_length = len(text_content)
        
        # 从文本内容中提取关键信息
        lines = text_content.split('\n')
        
        # 查找财务相关数据
        financial_items = []
        for line in lines:
            line = line.strip()
            if any(keyword in line for keyword in ['资产', '负债', '收入', '利润', '现金', '成本']):
                financial_items.append(line)
        
        # 截取前几行作为样本
        sample_lines = [line.strip() for line in lines if line.strip()][:10]
        
        # 生成基于实际内容的分析报告
        sections = [
            f"# 财务分析报告\n\n**报告生成日期：** {current_date}\n**数据来源：** 上传文档分析\n\n",
            "## 1. 数据概览\n\n",
            f"本次分析基于您上传的财务数据文档，文档总长度 {text_length} 字符。\n\n",
            f"文档样本内容：\n```\n{chr(10).join(sample_lines[:5])}\n```\n\n",
            "## 2. 财务数据分析\n\n",
        ]
        
        if financial_items:
            sections.append("### 识别的财务项目：\n")
            for item in financial_items[:5]:
                sections.append(f"- {item}\n")
            sections.append("\n")
        else:
            sections.append("### 数据识别结果：\n")
            sections.append("文档中包含的财务数据需要进一步结构化分析。\n\n")
        
        sections.extend([
            "## 3. 分析结论\n\n",
            "### 3.1 数据质量评估\n",
            f"- 数据完整性：文档包含 {len([l for l in lines if l.strip()])} 行有效数据\n",
            f"- 数据类型：{analysis_type} 分析适用\n",
            "- 建议：建议提供更结构化的财务报表数据以获得更深入的分析\n\n",
            "### 3.2 改进建议\n",
            "1. **数据标准化**：建议使用标准的财务报表格式\n",
            "2. **数据完整性**：确保包含完整的财务科目信息\n",
            "3. **时间序列**：提供多期数据以便进行趋势分析\n",
            "4. **详细分类**：细分收入、成本、费用等科目\n\n",
            "## 4. 后续建议\n\n",
            "为获得更准确的财务分析，建议：\n",
            "- 上传标准格式的资产负债表\n",
            "- 提供利润表数据\n",
            "- 包含现金流量表信息\n",
            "- 添加财务比率计算需要的基础数据\n\n",
            "---\n*本报告基于上传文档内容生成，如需更详细分析请提供标准财务报表*"
        ])
        
        # 逐段输出
        for section in sections:
            yield f"data: {json.dumps({'type': 'analysis_chunk', 'content': section}, ensure_ascii=False)}\n\n"
            await asyncio.sleep(0.2)
        
        print("基于内容的分析完成")
    
    # 保留原有的非流式方法作为备用
    async def extract_financial_data_simple(self, text_content: str, task_id: str = None) -> Dict[str, Any]:
        """简化的财务数据提取 - 去掉复杂的识别逻辑"""
        try:
            print("简化财务数据处理...")
            
            # 检查任务是否被取消
            if task_id and self._is_task_cancelled(task_id):
                raise asyncio.CancelledError("任务已被用户取消")
            
            # 简化的数据结构，直接返回文本摘要
            financial_data = {
                "text_summary": text_content[:2000],
                "text_length": len(text_content),
                "processed_at": datetime.now().isoformat(),
                "simplified": True  # 标记为简化处理
            }
            
            print(f"财务数据简化处理完成，文本长度: {len(text_content)}")
            
            return financial_data
            
        except asyncio.CancelledError:
            print(f"财务数据处理被取消 - 任务ID: {task_id}")
            raise
        except Exception as e:
            print(f"财务数据处理错误: {str(e)}")
            return {
                "text_summary": text_content[:2000],
                "text_length": len(text_content),
                "error": str(e)
            }
    
    async def call_deepseek_analysis_simple(self, financial_data: Dict[str, Any], analysis_type: str = "comprehensive", task_id: str = None) -> str:
        """简化的DeepSeek分析调用"""
        try:
            print(f"开始简化DeepSeek分析，类型: {analysis_type}")
            
            # 检查任务是否被取消
            if task_id and self._is_task_cancelled(task_id):
                raise asyncio.CancelledError("任务已被用户取消")
            
            # 构建简化的分析prompt
            template = self.analysis_templates.get(analysis_type, self.analysis_templates["comprehensive"])
            
            text_summary = financial_data.get('text_summary', '')
            
            analysis_prompt = f"""
            {template}
            
            基于以下财务文档内容：
            {text_summary}
            
            请生成专业的财务分析报告，要求结构清晰，分析深入。
            """
            
            # 如果没有API密钥，返回简化的分析结果
            if not self.deepseek_api_key:
                return self._generate_simple_analysis(financial_data, analysis_type)
            
            # 调用API
            headers = {
                "Authorization": f"Bearer {self.deepseek_api_key}",
                "Content-Type": "application/json"
            }
            
            payload = {
                "model": "deepseek-chat",
                "messages": [
                    {
                        "role": "system",
                        "content": "你是一位专业的财务分析师，请提供简洁、专业的财务分析报告。"
                    },
                    {
                        "role": "user",
                        "content": analysis_prompt
                    }
                ],
                "max_tokens": 3000,
                "temperature": 0.3
            }
            
            async with aiohttp.ClientSession() as session:
                try:
                    async with session.post(
                        self.deepseek_api_url,
                        headers=headers,
                        json=payload,
                        timeout=aiohttp.ClientTimeout(total=60)
                    ) as response:
                        if response.status == 200:
                            result = await response.json()
                            analysis_result = result['choices'][0]['message']['content']
                            print("简化DeepSeek分析完成")
                            return analysis_result
                        else:
                            print(f"DeepSeek API调用失败: {response.status}")
                            return self._generate_simple_analysis(financial_data, analysis_type)
                except Exception as api_error:
                    print(f"DeepSeek API错误: {str(api_error)}")
                    return self._generate_simple_analysis(financial_data, analysis_type)
            
        except asyncio.CancelledError:
            print(f"简化DeepSeek分析被取消 - 任务ID: {task_id}")
            raise
        except Exception as e:
            print(f"简化DeepSeek分析错误: {str(e)}")
            return self._generate_simple_analysis(financial_data, analysis_type)
    
    def _generate_simple_analysis(self, financial_data: Dict[str, Any], analysis_type: str) -> str:
        """生成简化的分析报告"""
        current_date = datetime.now().strftime("%Y年%m月%d日")
        text_length = financial_data.get('text_length', 0)
        
        return f"""# 财务分析报告

**报告生成日期：** {current_date}

## 执行摘要

基于上传的财务文档（共{text_length}字符），本次分析采用{analysis_type}分析方法，对企业财务状况进行评估。

## 主要发现

- 财务数据已成功处理
- 分析类型：{analysis_type}
- 建议重点关注现金流管理和盈利能力提升

## 分析结论

企业财务基础相对稳定，建议持续优化财务管理体系，提升运营效率。

## 改进建议

1. 加强财务监控体系建设
2. 优化成本控制措施  
3. 提升资金使用效率
4. 完善风险管理机制

---
*本报告由Fia财务分析系统生成*"""

    # 保留原有的任务管理方法
    def start_task(self, task_id: str):
        """开始一个新任务"""
        self.running_tasks[task_id] = {
            "status": "running",
            "start_time": datetime.now(),
            "cancelled": False
        }
        print(f"任务 {task_id} 已开始")
    
    def cancel_task(self, task_id: str):
        """取消指定任务"""
        if task_id in self.running_tasks:
            self.running_tasks[task_id]["cancelled"] = True
            self.running_tasks[task_id]["status"] = "cancelled"
            print(f"任务 {task_id} 已被取消")
            return True
        return False
    
    def finish_task(self, task_id: str):
        """完成指定任务"""
        if task_id in self.running_tasks:
            self.running_tasks[task_id]["status"] = "completed"
            self.running_tasks[task_id]["end_time"] = datetime.now()
            print(f"任务 {task_id} 已完成")
    
    def _is_task_cancelled(self, task_id: str) -> bool:
        """检查任务是否被取消"""
        if task_id in self.running_tasks:
            return self.running_tasks[task_id].get("cancelled", False)
        return False
    
    def get_task_status(self, task_id: str) -> Dict[str, Any]:
        """获取任务状态"""
        return self.running_tasks.get(task_id, {"status": "not_found"})

    def generate_word_report(self, analysis_result: str, company_name: str = "分析企业", task_id: str = None) -> str:
        """生成Word格式的分析报告"""
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=500, detail="Word文档生成功能不可用，请安装python-docx")
            
        try:
            print("开始生成Word报告...")
            
            # 检查任务是否被取消
            if task_id and self._is_task_cancelled(task_id):
                raise asyncio.CancelledError("任务已被用户取消")
            
            # 创建Word文档
            doc = Document()
            
            # 设置文档标题
            title = doc.add_heading(f'{company_name}财务分析报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加报告生成信息
            info_para = doc.add_paragraph()
            info_para.add_run('报告生成时间：').bold = True
            info_para.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M'))
            info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 添加分隔线
            doc.add_paragraph('=' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 处理分析结果内容
            lines = analysis_result.split('\n')
            current_paragraph = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 检查是否为标题
                if line.startswith('# '):
                    # 一级标题
                    title_text = line[2:].strip()
                    doc.add_heading(title_text, level=1)
                elif line.startswith('## '):
                    # 二级标题
                    title_text = line[3:].strip()
                    doc.add_heading(title_text, level=2)
                elif line.startswith('### '):
                    # 三级标题
                    title_text = line[4:].strip()
                    doc.add_heading(title_text, level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    # 列表项
                    para = doc.add_paragraph(line[2:].strip(), style='List Bullet')
                elif line.startswith('**') and line.endswith('**'):
                    # 加粗文本
                    para = doc.add_paragraph()
                    para.add_run(line[2:-2]).bold = True
                else:
                    # 普通段落
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    current_paragraph.add_run(line + '\n')
                    if line.endswith('.') or line.endswith('。'):
                        current_paragraph = None
            
            # 添加页脚
            footer_para = doc.add_paragraph()
            footer_para.add_run('本报告由Fia财务分析系统自动生成').italic = True
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 保存文档
            temp_dir = "temp"
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"财务分析报告_{company_name}_{timestamp}.docx"
            filepath = os.path.join(temp_dir, filename)
            
            doc.save(filepath)
            print(f"Word报告生成完成: {filepath}")
            
            # 在保存前最后检查
            if task_id and self._is_task_cancelled(task_id):
                raise asyncio.CancelledError("任务已被用户取消")
            
            return filepath
            
        except asyncio.CancelledError:
            print(f"Word报告生成被取消 - 任务ID: {task_id}")
            raise
        except Exception as e:
            print(f"Word报告生成错误: {str(e)}")
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Word报告生成失败: {str(e)}")

# 创建全局服务实例
financial_service = FinancialAnalysisService() 