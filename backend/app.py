"""财务分析后端服务"""
import os
import sys
import asyncio
import tempfile
from typing import Dict, List, Optional
import numpy as np
import uvicorn
import pandas as pd
from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from fastapi.responses import FileResponse, StreamingResponse
import json
import uuid
from collections import Counter

# 添加必要的导入
import faiss
from sklearn.feature_extraction.text import TfidfVectorizer
import jieba
import re
from config import BASE_CONFIG, VECTOR_CONFIG, FILE_CONFIG, ensure_temp_dir

# 确保临时目录存在
ensure_temp_dir()

# 创建FastAPI应用
app = FastAPI(title="财务分析API", version="1.0.0")

# 添加CORS中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 全局变量
vector_dimension = VECTOR_CONFIG["VECTOR_DIMENSION"]
index = faiss.IndexFlatL2(vector_dimension)
document_store = {}
tfidf_vectorizer = TfidfVectorizer(max_features=vector_dimension, stop_words=None)
tfidf_fitted = False

# Pydantic模型
class QuestionRequest(BaseModel):
    text: str

class FinancialAnalysisRequest(BaseModel):
    analysis_type: str = "comprehensive"  # comprehensive, profitability, liquidity, efficiency
    company_name: Optional[str] = "分析企业"

class FinancialReportRequest(BaseModel):
    doc_ids: Optional[List[int]] = None  # 要分析的文档ID列表，None表示使用所有文档
    analysis_type: str = "comprehensive"
    company_name: Optional[str] = "分析企业"

class TaskCancelRequest(BaseModel):
    task_id: str

class WordOnlyRequest(BaseModel):
    analysis_result: str
    analysis_type: str = "comprehensive"
    company_name: Optional[str] = "分析企业"

def simple_text_vectorize(text: str) -> np.ndarray:
    """改进的文本向量化方法"""
    global tfidf_vectorizer, tfidf_fitted
    
    # 如果TF-IDF未拟合，先用当前文本进行简单初始化
    if not tfidf_fitted:
        print("TF-IDF未初始化，使用当前文本进行初始化")
        tfidf_vectorizer.fit([text])
        tfidf_fitted = True
    
    # 预处理文本
    def preprocess_text(text):
        # 移除多余空白
        text = re.sub(r'\s+', ' ', text)
        # 保留数字和小数点
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,\-]', '', text)
        return text
    
    # 处理文本
    processed_text = preprocess_text(text)
    
    # 中文分词
    words = list(jieba.cut(processed_text))
    
    # 对数字进行特殊处理
    processed_words = []
    for word in words:
        if re.match(r'^[\d.,\-]+$', word):
            # 保持数字的完整性
            processed_words.append(f"NUM_{word}")
        else:
            processed_words.append(word)
    
    # 合并处理后的词
    processed_text = ' '.join(processed_words)
    
    # 使用TF-IDF向量化
    vector = tfidf_vectorizer.transform([processed_text]).toarray().flatten()
    
    # 确保向量维度正确
    if len(vector) < vector_dimension:
        padded_vector = np.zeros(vector_dimension)
        padded_vector[:len(vector)] = vector
        return padded_vector.astype('float32')
    else:
        return vector[:vector_dimension].astype('float32')

def extract_text_from_pdf(file_path: str) -> str:
    """改进的PDF文本提取方法，支持表格和图片"""
    try:
        import fitz  # PyMuPDF
        print(f"\n=== PDF处理开始 ===")
        print(f"处理文件: {file_path}")
        
        # 打开PDF文件
        doc = fitz.open(file_path)
        print(f"PDF页数: {len(doc)}")
        
        all_text = []
        
        for page_num, page in enumerate(doc):
            print(f"\n--- 处理第 {page_num + 1} 页 ---")
            
            # 方法1: 直接提取文本
            text = page.get_text()
            print(f"直接提取文本长度: {len(text)}")
            
            # 方法2: 提取表格
            tables = []
            try:
                # 尝试提取表格
                table_data = page.find_tables()
                for table in table_data:
                    table_text = []
                    try:
                        table_content = table.extract()
                        for row in table_content:
                            if row and any(cell for cell in row):  # 过滤空行
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                table_text.append(row_text)
                        if table_text:
                            tables.append("表格内容:\n" + "\n".join(table_text))
                    except Exception as e:
                        print(f"表格提取错误: {e}")
                        continue
                print(f"提取到 {len(tables)} 个表格")
            except Exception as e:
                print(f"表格查找错误: {e}")
            
            # 方法3: 基于文本块的提取
            text_blocks = []
            try:
                blocks = page.get_text("dict")
                for block in blocks["blocks"]:
                    if "lines" in block:
                        block_text = ""
                        for line in block["lines"]:
                            for span in line["spans"]:
                                if span.get("text", "").strip():
                                    block_text += span["text"] + " "
                        if block_text.strip():
                            text_blocks.append(block_text.strip())
                print(f"提取到 {len(text_blocks)} 个文本块")
            except Exception as e:
                print(f"文本块提取错误: {e}")
            
            # 合并所有提取的内容
            page_content = []
            
            # 添加直接提取的文本
            if text.strip():
                page_content.append(f"页面文本:\n{text.strip()}")
            
            # 添加表格内容
            if tables:
                page_content.extend(tables)
            
            # 添加文本块内容（如果直接提取为空）
            if not text.strip() and text_blocks:
                page_content.append("文本块内容:\n" + "\n".join(text_blocks))
            
            # 如果仍然没有内容，尝试OCR提取（基础方法）
            if not page_content:
                try:
                    # 尝试获取页面图像并转换为文本（简单方法）
                    pix = page.get_pixmap()
                    if pix.width > 0 and pix.height > 0:
                        page_content.append("注意：此页面可能包含图片或扫描内容，需要OCR处理")
                        print("页面可能包含图片内容")
                except Exception as e:
                    print(f"图像处理错误: {e}")
            
            if page_content:
                page_text = f"\n=== 第 {page_num + 1} 页 ===\n" + "\n\n".join(page_content)
                all_text.append(page_text)
                print(f"第 {page_num + 1} 页最终文本长度: {len(page_text)}")
            else:
                print(f"第 {page_num + 1} 页未提取到内容")
        
        doc.close()
        
        # 合并所有页面文本
        final_text = "\n\n".join(all_text)
        print(f"\n总文本长度: {len(final_text)}")
        print(f"文本预览: {final_text[:200]}...")
        print("=== PDF处理完成 ===\n")
        
        return final_text
        
    except Exception as e:
        print(f"\n=== PDF处理错误 ===")
        print(f"错误信息: {str(e)}")
        print(f"文件路径: {file_path}")
        try:
            print(f"文件大小: {os.path.getsize(file_path)} 字节")
        except:
            print("无法获取文件大小")
        print("=== 错误详情结束 ===\n")
        
        # 即使出错也尝试返回基础提取结果
        try:
            import fitz
            doc = fitz.open(file_path)
            basic_text = ""
            for page in doc:
                basic_text += page.get_text() + "\n"
            doc.close()
            if basic_text.strip():
                return f"基础提取结果:\n{basic_text.strip()}"
        except:
            pass
            
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """从DOCX文件提取文本，包括表格内容"""
    try:
        from docx import Document
        print(f"开始处理Word文档: {os.path.basename(file_path)}")
        
        doc = Document(file_path)
        all_content = []
        
        # 处理文档中的段落和表格
        for element in doc.element.body:
            if element.tag.endswith('p'):  # 段落
                # 找到对应的段落对象
                for para in doc.paragraphs:
                    if para._element == element:
                        if para.text.strip():
                            all_content.append(para.text.strip())
                        break
            
            elif element.tag.endswith('tbl'):  # 表格
                # 找到对应的表格对象
                for table in doc.tables:
                    if table._element == element:
                        table_content = []
                        table_content.append("=== 表格内容 ===")
                        
                        # 处理表格头部
                        if table.rows:
                            header_row = table.rows[0]
                            headers = []
                            for cell in header_row.cells:
                                headers.append(cell.text.strip() if cell.text.strip() else "")
                            if any(headers):
                                table_content.append("表头: " + " | ".join(headers))
                        
                        # 处理表格数据行
                        for row_idx, row in enumerate(table.rows[1:], 1):
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                row_data.append(cell_text if cell_text else "")
                            if any(row_data):
                                # 如果有表头，按列名组织数据
                                if headers and len(row_data) == len(headers):
                                    formatted_row = []
                                    for header, data in zip(headers, row_data):
                                        if data:
                                            formatted_row.append(f"{header}: {data}")
                                    if formatted_row:
                                        table_content.append(" | ".join(formatted_row))
                                else:
                                    table_content.append(" | ".join(row_data))
                        
                        if len(table_content) > 1:  # 有实际内容
                            all_content.append("\n".join(table_content))
                        break
        
        final_text = "\n\n".join(all_content)
        print(f"Word文档处理完成，提取文本长度: {len(final_text)}")
        print(f"包含段落数: {len([c for c in all_content if not c.startswith('=== 表格内容')])}")
        print(f"包含表格数: {len([c for c in all_content if c.startswith('=== 表格内容')])}")
        
        return final_text
        
    except Exception as e:
        print(f"Word文档处理错误: {str(e)}")
        # 回退到基础处理
        try:
            from docx import Document
            doc = Document(file_path)
            basic_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            return f"Word文档基础提取:\n{basic_text}"
        except:
            return f"DOCX解析错误: {str(e)}"

def extract_text_from_excel(file_path: str) -> str:
    """从Excel文件提取文本，保留表格结构"""
    try:
        # 读取所有工作表
        xl_file = pd.ExcelFile(file_path)
        all_text = []
        
        for sheet_name in xl_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # 添加工作表名称
            sheet_text = [f"工作表: {sheet_name}"]
            
            # 添加列名
            headers = df.columns.tolist()
            sheet_text.append("列标题: " + ", ".join(str(h) for h in headers))
            
            # 处理每一行数据
            for idx, row in df.iterrows():
                row_text = []
                for col in headers:
                    value = row[col]
                    if pd.notna(value):  # 只处理非空值
                        row_text.append(f"{col}: {value}")
                if row_text:
                    sheet_text.append(" | ".join(row_text))
            
            # 合并工作表文本
            all_text.append("\n".join(sheet_text))
        
        final_text = "\n\n".join(all_text)
        print(f"Excel文件 {os.path.basename(file_path)} 提取文本长度: {len(final_text)}")
        return final_text
        
    except Exception as e:
        print(f"Excel处理错误 ({os.path.basename(file_path)}): {str(e)}")
        raise Exception(f"Excel文件处理失败: {str(e)}")

def extract_text_from_csv(file_path: str) -> str:
    """从CSV文件提取文本，保留表格结构"""
    try:
        print(f"开始处理CSV文件: {os.path.basename(file_path)}")
        
        # 尝试不同的编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        df = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                used_encoding = encoding
                print(f"成功使用编码 {encoding} 读取CSV文件")
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"使用编码 {encoding} 读取失败: {e}")
                continue
        
        if df is None:
            raise Exception("无法使用任何编码读取CSV文件")
        
        # 构建结构化文本
        csv_content = []
        csv_content.append("=== CSV数据 ===")
        
        # 添加基本信息
        csv_content.append(f"数据行数: {len(df)}")
        csv_content.append(f"数据列数: {len(df.columns)}")
        csv_content.append(f"使用编码: {used_encoding}")
        
        # 添加列名
        headers = df.columns.tolist()
        csv_content.append("列名: " + ", ".join(str(h) for h in headers))
        
        # 添加数据行（限制显示行数以避免过长）
        max_rows = min(1000, len(df))  # 最多显示1000行
        csv_content.append(f"\n数据内容（前{max_rows}行）:")
        
        for idx, row in df.head(max_rows).iterrows():
            row_text = []
            for col in headers:
                value = row[col]
                if pd.notna(value):
                    row_text.append(f"{col}: {value}")
            if row_text:
                csv_content.append(" | ".join(row_text))
        
        if len(df) > max_rows:
            csv_content.append(f"... (还有 {len(df) - max_rows} 行数据)")
        
        final_text = "\n".join(csv_content)
        print(f"CSV文件处理完成，提取文本长度: {len(final_text)}")
        
        return final_text
        
    except Exception as e:
        print(f"CSV处理错误: {str(e)}")
        # 回退到基础文本处理
        try:
            for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    return f"CSV基础文本提取（编码: {encoding}）:\n{content}"
                except:
                    continue
            raise Exception("所有编码尝试失败")
        except:
            return f"CSV解析错误: {str(e)}"

def clean_text(text: str) -> str:
    """清理文本的简化版本"""
    if not text:
        return ""
        
    # 只做基本的清理
    text = re.sub(r'\s+', ' ', text)  # 合并多个空白字符
    text = text.strip()  # 移除首尾空白
    return text

def split_into_chunks(text: str, chunk_size: int = 1000) -> List[str]:
    """改进的文本分块函数"""
    if not text:
        return []
        
    # 如果文本较短，直接返回
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    
    # 按段落分割
    paragraphs = text.split('\n\n')
    current_chunk = ''
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:
            continue
            
        # 如果当前段落加上现有块不超过限制
        if len(current_chunk) + len(paragraph) + 2 <= chunk_size:
            if current_chunk:
                current_chunk += '\n\n' + paragraph
            else:
                current_chunk = paragraph
        else:
            # 保存当前块
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            # 如果单个段落太长，按句子分割
            if len(paragraph) > chunk_size:
                sentences = paragraph.replace('。', '。\n').split('\n')
                temp_chunk = ''
                for sentence in sentences:
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                        
                    if len(temp_chunk) + len(sentence) + 1 <= chunk_size:
                        if temp_chunk:
                            temp_chunk += sentence
                        else:
                            temp_chunk = sentence
                    else:
                        if temp_chunk.strip():
                            chunks.append(temp_chunk.strip())
                        temp_chunk = sentence
                
                if temp_chunk.strip():
                    current_chunk = temp_chunk
                else:
                    current_chunk = ''
            else:
                current_chunk = paragraph
    
    # 添加最后一个块
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # 过滤过短的块
    filtered_chunks = [chunk for chunk in chunks if len(chunk.strip()) > 20]
    
    # 如果没有有效块，返回原文本
    if not filtered_chunks:
        return [text]
    
    return filtered_chunks

def fit_tfidf_vectorizer(all_texts: List[str]):
    """拟合TF-IDF向量化器"""
    global tfidf_vectorizer, tfidf_fitted
    
    # 对所有文本进行中文分词
    processed_texts = []
    for text in all_texts:
        words = list(jieba.cut(text))
        processed_texts.append(' '.join(words))
    
    # 拟合TF-IDF向量化器
    tfidf_vectorizer.fit(processed_texts)
    tfidf_fitted = True
    print("TF-IDF向量化器拟合完成")

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """文件上传处理 - 简化版"""
    print(f"\n=== 开始处理上传文件 ===")
    print(f"文件名: {file.filename}")
    
    try:
        # 基本检查
        if not file.filename:
            result = {
                "success": False,
                "message": "文件名不能为空"
            }
            print(f"返回结果: {result}")
            return result
        
        # 获取文件扩展名
        file_ext = os.path.splitext(file.filename)[1].lower()
        print(f"文件类型: {file_ext}")
        
        # 检查支持的文件类型
        supported_types = ['.pdf', '.xlsx', '.xls', '.docx', '.txt', '.csv']
        if file_ext not in supported_types:
            result = {
                "success": False,
                "message": f"不支持的文件类型: {file_ext}。支持的格式：{', '.join(supported_types)}"
            }
            print(f"返回结果: {result}")
            return result
        
        # 创建临时目录
        temp_dir = "temp"
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
            print(f"创建临时目录: {temp_dir}")
        
        # 保存文件
        temp_path = os.path.join(temp_dir, f"temp_{file.filename}")
        content = await file.read()
        
        if len(content) == 0:
            result = {
                "success": False,
                "message": "上传的文件为空"
            }
            print(f"返回结果: {result}")
            return result
        
        with open(temp_path, "wb") as f:
            f.write(content)
        
        print(f"文件已保存到: {temp_path}")
        print(f"文件大小: {len(content)} 字节")
        
        # 提取文本
        text = ""
        try:
            if file_ext == '.txt':
                # 纯文本文件处理
                for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                    try:
                        with open(temp_path, 'r', encoding=encoding) as f:
                            text = f.read()
                        print(f"成功使用编码 {encoding} 读取TXT文件")
                        break
                    except UnicodeDecodeError:
                        continue
                if not text:
                    result = {
                        "success": False,
                        "message": "无法解码TXT文件，请检查文件编码"
                    }
                    print(f"返回结果: {result}")
                    return result
            elif file_ext == '.csv':
                # CSV文件专门处理
                print(f"开始处理CSV文件: {file.filename}")
                text = extract_text_from_csv(temp_path)
                print(f"CSV文件处理完成，提取文本长度: {len(text)}")
            elif file_ext in ['.xlsx', '.xls']:
                # Excel文件支持
                print(f"开始处理Excel文件: {file.filename}")
                text = extract_text_from_excel(temp_path)
                print(f"Excel文件处理完成，提取文本长度: {len(text)}")
            elif file_ext == '.docx':
                # Word文档支持（包含表格）
                print(f"开始处理Word文档: {file.filename}")
                text = extract_text_from_docx(temp_path)
                print(f"Word文档处理完成，提取文本长度: {len(text)}")
            elif file_ext == '.pdf':
                # PDF文件支持
                print(f"开始处理PDF文件: {file.filename}")
                text = extract_text_from_pdf(temp_path)
                print(f"PDF文件处理完成，提取文本长度: {len(text)}")
            else:
                # 其他文件类型暂不支持
                result = {
                    "success": False,
                    "message": f"暂时不支持 {file_ext} 文件类型的处理"
                }
                print(f"返回结果: {result}")
                return result
                
        except Exception as extract_error:
            print(f"文本提取错误: {str(extract_error)}")
            result = {
                "success": False,
                "message": f"文本提取失败: {str(extract_error)}"
            }
            print(f"返回结果: {result}")
            return result
        
        if not text or len(text.strip()) < 10:
            # 对于PDF文件，提供更详细的错误信息
            if file_ext == '.pdf':
                result = {
                    "success": False,
                    "message": f"PDF文件 {file.filename} 文本提取失败。可能原因：\n"
                             f"1. PDF包含扫描图片，需要OCR处理\n"
                             f"2. PDF格式特殊或已加密\n"
                             f"3. PDF主要包含图形或表格，文本内容较少\n"
                             f"提取的文本长度: {len(text.strip()) if text else 0}"
                }
            elif file_ext == '.docx':
                result = {
                    "success": False,
                    "message": f"Word文档 {file.filename} 文本提取失败。可能原因：\n"
                             f"1. Word文档格式特殊或已加密\n"
                             f"2. Word文档主要包含图形或表格，文本内容较少\n"
                             f"提取的文本长度: {len(text.strip()) if text else 0}"
                }
            else:
                result = {
                    "success": False,
                    "message": f"无法从文件 {file.filename} 中提取有效文本内容，文本长度: {len(text.strip()) if text else 0}"
                }
            print(f"返回结果: {result}")
            return result
        
        print(f"成功提取文本，长度: {len(text)}")
        
        # 简化处理：直接存储文本，不进行向量化
        doc_id = len(document_store)
        document_store[doc_id] = {
            'filename': file.filename,
            'text': text,
            'chunks': [text],  # 简化：不分块
            'vector_ids': []
        }
        
        print(f"文档已存储，ID: {doc_id}")
        
        result = {
            "success": True,
            "message": f"成功处理文件：{file.filename}",
            "stats": {
                "text_length": len(text),
                "chunks_count": 1,
                "vector_dimension": 0,
                "vectors_created": 0
            }
        }
        
        print(f"返回结果: {result}")
        return result
            
    except Exception as e:
        print(f"\n=== 处理错误 ===")
        print(f"错误信息: {str(e)}")
        import traceback
        traceback.print_exc()
        
        result = {
            "success": False,
            "message": f"处理文件失败: {str(e)}"
        }
        print(f"返回错误结果: {result}")
        return result
        
    finally:
        # 清理临时文件
        if 'temp_path' in locals() and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
                print(f"已删除临时文件: {temp_path}")
            except Exception as e:
                print(f"删除临时文件失败: {str(e)}")
        print("=== 文件处理结束 ===\n")

@app.post("/ask")
async def ask_question(question: QuestionRequest):
    """处理问答请求"""
    print(f"\n=== 问答请求开始 ===")
    print(f"问题: {question.text}")
    print(f"当前文档数量: {len(document_store)}")
    print(f"文档存储详情: {list(document_store.keys())}")
    print(f"FAISS索引向量数: {index.ntotal}")
    
    # 详细检查文档状态
    if len(document_store) == 0:
        print("❌ 没有文档，document_store为空")
        print(f"document_store内容: {document_store}")
        return {
            "success": False,
            "message": f"没有上传任何文档，请先上传文档。当前文档数量: {len(document_store)}"
        }
    
    print(f"✅ 找到 {len(document_store)} 个文档")
    for doc_id, doc_info in document_store.items():
        print(f"  文档 {doc_id}: {doc_info['filename']}, 文本长度: {len(doc_info['text'])}")
    
    try:
        print("开始处理问答...")
        
        # 使用更高效的文本搜索算法
        relevant_docs = improved_text_search(question.text, list(document_store.values()))
        
        # 如果没有找到相关文档，返回所有文档的片段
        if not relevant_docs:
            print("没有找到匹配的文档，返回所有文档")
            relevant_docs = [{
                'text': doc['text'][:8000],  # 大幅增加到8000字符
                'filename': doc['filename'],
                'similarity': 0.1
            } for doc in document_store.values()]
        
        # 构建上下文 - 符合前端API期望的格式
        context_parts = []
        relevant_chunks = []
        
        for i, doc in enumerate(relevant_docs, 1):
            context_parts.append(f"[文件: {doc['filename']} | 相似度: {doc['similarity']:.2f}]\n{doc['text'][:4000]}...")  # 大幅增加到4000字符
            relevant_chunks.append({
                'text': doc['text'][:2000],  # 大幅增加到2000字符
                'filename': doc['filename'],
                'similarity': doc['similarity']
            })
        
        context = "\n\n".join(context_parts)
        
        # 返回符合前端API期望的格式
        result = {
            "success": True,
            "answer": {
                "context": context,
                "relevant_chunks": relevant_chunks
            }
        }
        
        print(f"✅ 问答处理完成，返回结果格式正确")
        return result
        
    except Exception as e:
        print(f"❌ 问答处理错误: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"处理问题失败: {str(e)}"
        }
    finally:
        print("=== 问答请求结束 ===\n")

@app.get("/status")
async def get_status():
    """获取系统状态"""
    return {
        "success": True,
        "stats": {
            "total_documents": len(document_store),
            "total_vectors": index.ntotal,
            "tfidf_fitted": tfidf_fitted,
            "files": [
                {
                    "filename": info["filename"],
                    "chunks": len(info["chunks"])
                }
                for info in document_store.values()
            ]
        }
    }

@app.delete("/documents")
async def clear_documents():
    """清空所有文档"""
    global document_store, index, tfidf_fitted
    
    # 清空文档存储
    document_store.clear()
    
    # 重新创建索引
    index = faiss.IndexFlatL2(vector_dimension)
    
    # 重置TF-IDF
    tfidf_fitted = False
    
    return {
        "success": True,
        "message": "所有文档已清空"
    }

@app.post("/debug/reset")
async def debug_reset():
    """调试：重置所有状态"""
    global document_store, index, tfidf_fitted, tfidf_vectorizer
    
    # 清空所有状态
    document_store.clear()
    index = faiss.IndexFlatL2(vector_dimension)
    tfidf_fitted = False
    tfidf_vectorizer = TfidfVectorizer(max_features=vector_dimension, stop_words=None)
    
    print("=== 调试重置完成 ===")
    
    return {
        "success": True,
        "message": "系统状态已重置",
        "stats": {
            "total_documents": len(document_store),
            "total_vectors": index.ntotal,
            "tfidf_fitted": tfidf_fitted
        }
    }

@app.get("/")
async def root():
    """根路径"""
    return {
        "message": "财务分析后端服务",
        "version": "1.0.0",
        "status": "运行中"
    }

@app.post("/financial-analysis")
async def generate_financial_analysis_fast(request: FinancialAnalysisRequest):
    """快速财务分析报告生成（简化版本）"""
    task_id = str(uuid.uuid4())
    
    try:
        print(f"\n=== 开始快速财务分析 ===")
        print(f"任务ID: {task_id}")
        print(f"分析类型: {request.analysis_type}")
        print(f"公司名称: {request.company_name}")
        print(f"当前文档数量: {len(document_store)}")
        
        # 检查是否有上传的文档
        if len(document_store) == 0:
            return {
                "success": False,
                "message": "请先上传财务数据文件（Excel、PDF等格式）"
            }
        
        # 收集所有文档的文本内容
        all_text = ""
        processed_files = []
        
        for doc_id, doc_info in document_store.items():
            all_text += f"\n\n=== 文件: {doc_info['filename']} ===\n"
            all_text += doc_info['text']
            processed_files.append(doc_info['filename'])
        
        print(f"处理文件: {processed_files}")
        print(f"总文本长度: {len(all_text)}")
        
        # 导入财务分析服务并启动任务
        from financial_analysis import financial_service
        financial_service.start_task(task_id)
        
        # 简化的财务数据处理
        financial_data = await financial_service.extract_financial_data_simple(all_text, task_id)
        
        # 调用简化的DeepSeek分析
        analysis_result = await financial_service.call_deepseek_analysis_simple(
            financial_data, 
            request.analysis_type,
            task_id
        )
        
        # 完成任务
        financial_service.finish_task(task_id)
        
        print("快速财务分析完成")
        
        return {
            "success": True,
            "message": "财务分析报告生成成功",
            "data": {
                "task_id": task_id,
                "analysis_result": analysis_result,
                "processed_files": processed_files,
                "analysis_type": request.analysis_type,
                "company_name": request.company_name,
                "processing_time": "优化版本，响应更快",
                "simplified": True
            }
        }
        
    except asyncio.CancelledError:
        return {
            "success": False,
            "message": "分析任务已被用户取消",
            "task_id": task_id,
            "cancelled": True
        }
    except Exception as e:
        financial_service.finish_task(task_id)
        print(f"快速财务分析错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"财务分析失败: {str(e)}"
        }

@app.post("/financial-report/generate")
async def generate_financial_report(request: FinancialReportRequest):
    """生成并返回Word格式的财务分析报告"""
    task_id = str(uuid.uuid4())
    
    try:
        print(f"\n=== 生成财务报告 ===")
        print(f"任务ID: {task_id}")
        print(f"文档IDs: {request.doc_ids}")
        print(f"分析类型: {request.analysis_type}")
        print(f"公司名称: {request.company_name}")
        
        # 导入财务分析服务
        try:
            from financial_analysis import financial_service
        except ImportError:
            return {
                "success": False,
                "message": "财务分析模块未正确安装，请检查依赖"
            }
        
        # 验证文档ID
        if not request.doc_ids:
            # 如果没有指定文档ID，使用所有已上传的文档
            if len(document_store) == 0:
                return {
                    "success": False,
                    "message": "请先上传财务数据文件"
                }
            request.doc_ids = list(document_store.keys())
        
        # 收集指定文档的文本
        all_text = ""
        processed_files = []
        
        for doc_id in request.doc_ids:
            if doc_id in document_store:
                doc_info = document_store[doc_id]
                all_text += f"\n\n=== 文件: {doc_info['filename']} ===\n"
                all_text += doc_info['text']
                processed_files.append(doc_info['filename'])
            else:
                print(f"文档ID {doc_id} 不存在")
        
        if not all_text:
            return {
                "success": False,
                "message": "未找到有效的财务数据"
            }
        
        print(f"处理文件: {processed_files}")
        
        # 提取财务数据
        financial_data = await financial_service.extract_financial_data(all_text, task_id)
        
        # 生成分析报告
        analysis_result = await financial_service.call_deepseek_analysis(
            financial_data,
            request.analysis_type,
            task_id
        )
        
        # 生成Word文档 - 移除await，因为generate_word_report是同步方法
        word_filepath = financial_service.generate_word_report(
            analysis_result,
            request.company_name or "分析企业",
            task_id
        )
        
        financial_service.finish_task(task_id)
        
        return {
            "success": True,
            "message": "财务报告生成成功",
            "data": {
                "task_id": task_id,
                "report_file": os.path.basename(word_filepath),
                "download_url": f"/api/financial-report/download/{os.path.basename(word_filepath)}",
                "processed_files": processed_files,
                "analysis_type": request.analysis_type,
                "company_name": request.company_name,
                "file_path": word_filepath
            }
        }
        
    except asyncio.CancelledError:
        return {
            "success": False,
            "message": "报告生成任务已被用户取消",
            "task_id": task_id,
            "cancelled": True
        }
    except Exception as e:
        financial_service.finish_task(task_id)
        print(f"财务报告生成错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"财务报告生成失败: {str(e)}"
        }

@app.get("/financial-report/download/{filename}")
async def download_financial_report(filename: str):
    """下载财务分析报告Word文档"""
    try:
        print(f"下载财务报告: {filename}")
        
        # 构建文件路径
        filepath = os.path.join("temp", filename)
        
        # 检查文件是否存在
        if not os.path.exists(filepath):
            raise HTTPException(status_code=404, detail="报告文件不存在或已过期")
        
        # 修复：处理中文文件名编码问题
        import urllib.parse
        
        # 对文件名进行URL编码
        encoded_filename = urllib.parse.quote(filename.encode('utf-8'), safe='')
        
        # 返回文件下载响应
        return FileResponse(
            path=filepath,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                # 使用RFC 6266标准支持UTF-8文件名
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
                "Cache-Control": "no-cache",
                "Pragma": "no-cache"
            }
        )
        
    except UnicodeEncodeError as e:
        print(f"文件名编码错误: {str(e)}")
        # 如果文件名编码失败，使用简化文件名
        simple_filename = "财务分析报告.docx"
        return FileResponse(
            path=filepath,
            filename=simple_filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                "Content-Disposition": f"attachment; filename={simple_filename}",
                "Cache-Control": "no-cache",
                "Pragma": "no-cache"
            }
        )
    except Exception as e:
        print(f"下载错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件下载失败: {str(e)}")

@app.get("/financial-analysis/templates")
async def get_analysis_templates():
    """获取可用的财务分析模板"""
    return {
        "success": True,
        "data": {
            "templates": [
                {
                    "id": "comprehensive",
                    "name": "综合财务分析",
                    "description": "全面分析企业财务状况，包括盈利能力、偿债能力、营运能力和成长能力"
                },
                {
                    "id": "profitability",
                    "name": "盈利能力分析",
                    "description": "专项分析企业的盈利能力和收入质量"
                },
                {
                    "id": "liquidity",
                    "name": "流动性分析",
                    "description": "专项分析企业的流动性和偿债能力"
                },
                {
                    "id": "efficiency",
                    "name": "营运效率分析",
                    "description": "专项分析企业的营运效率和资产利用率"
                }
            ]
        }
    }

@app.delete("/financial-report/cleanup")
async def cleanup_reports():
    """清理临时生成的报告文件"""
    try:
        temp_dir = "temp"
        cleaned_count = 0
        if os.path.exists(temp_dir):
            # 删除所有.docx文件
            for filename in os.listdir(temp_dir):
                if filename.endswith('.docx'):
                    filepath = os.path.join(temp_dir, filename)
                    try:
                        os.remove(filepath)
                        print(f"删除报告文件: {filename}")
                        cleaned_count += 1
                    except Exception as e:
                        print(f"删除文件失败 {filename}: {str(e)}")
        
        return {
            "success": True,
            "message": f"报告文件清理完成，共清理 {cleaned_count} 个文件"
        }
        
    except Exception as e:
        print(f"清理错误: {str(e)}")
        return {
            "success": False,
            "message": f"清理失败: {str(e)}"
        }

@app.get("/financial-analysis/status")
async def get_financial_analysis_status():
    """获取财务分析功能状态"""
    try:
        # 检查财务分析模块
        try:
            from financial_analysis import financial_service
            financial_module_available = True
        except ImportError:
            financial_module_available = False
        
        # 检查DeepSeek配置
        deepseek_configured = bool(os.getenv("DEEPSEEK_API_KEY"))
        
        # 检查Word生成功能
        try:
            from docx import Document
            word_generation_available = True
        except ImportError:
            word_generation_available = False
        
        return {
            "success": True,
            "data": {
                "financial_module_available": financial_module_available,
                "deepseek_configured": deepseek_configured,
                "word_generation_available": word_generation_available,
                "uploaded_documents": len(document_store),
                "temp_reports": len([f for f in os.listdir("temp") if f.endswith('.docx')]) if os.path.exists("temp") else 0
            }
        }
        
    except Exception as e:
        return {
            "success": False,
            "message": f"状态检查失败: {str(e)}"
        }

@app.post("/financial-analysis/cancel")
async def cancel_financial_analysis(request: TaskCancelRequest):
    """取消财务分析任务"""
    try:
        print(f"取消财务分析任务: {request.task_id}")
        
        from financial_analysis import financial_service
        cancelled = financial_service.cancel_task(request.task_id)
        
        if cancelled:
            return {
                "success": True,
                "message": f"任务 {request.task_id} 已成功取消"
            }
        else:
            return {
                "success": False,
                "message": f"任务 {request.task_id} 不存在或已完成"
            }
            
    except Exception as e:
        return {
            "success": False,
            "message": f"取消任务失败: {str(e)}"
        }

@app.get("/financial-analysis/task-status/{task_id}")
async def get_task_status(task_id: str):
    """获取任务状态"""
    try:
        from financial_analysis import financial_service
        status = financial_service.get_task_status(task_id)
        
        return {
            "success": True,
            "data": {
                "task_id": task_id,
                "status": status
            }
        }
        
    except Exception as e:
        return {
            "success": False,
            "message": f"获取任务状态失败: {str(e)}"
        }

@app.post("/financial-report/generate-word-only")
async def generate_word_only(request: WordOnlyRequest):
    """基于已有分析结果生成Word文档"""
    task_id = str(uuid.uuid4())
    
    try:
        print(f"\n=== 生成Word文档（基于已有分析结果） ===")
        print(f"任务ID: {task_id}")
        print(f"分析类型: {request.analysis_type}")
        print(f"公司名称: {request.company_name}")
        
        # 导入财务分析服务
        try:
            from financial_analysis import financial_service
        except ImportError:
            return {
                "success": False,
                "message": "财务分析模块未正确安装，请检查依赖"
            }
        
        # 直接使用传入的分析结果生成Word文档
        word_filepath = financial_service.generate_word_report(
            request.analysis_result,
            request.company_name or "分析企业",
            task_id
        )
        
        return {
            "success": True,
            "message": "Word文档生成成功",
            "data": {
                "task_id": task_id,
                "report_file": os.path.basename(word_filepath),
                "download_url": f"/api/financial-report/download/{os.path.basename(word_filepath)}",
                "analysis_type": request.analysis_type,
                "company_name": request.company_name,
                "file_path": word_filepath
            }
        }
        
    except Exception as e:
        print(f"Word文档生成错误: {str(e)}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "message": f"Word文档生成失败: {str(e)}"
        }

def improved_text_search(question: str, docs: list, max_results: int = 3):
    """改进的文本搜索算法"""
    question_words = set(re.findall(r'\b\w{2,}\b', question.lower()))
    
    scored_docs = []
    for doc in docs:
        doc_words = set(re.findall(r'\b\w{2,}\b', doc['text'].lower()))
        
        # 计算词汇重叠度
        overlap = len(question_words & doc_words)
        total_words = len(question_words | doc_words)
        
        if total_words > 0:
            similarity = overlap / total_words
            scored_docs.append({
                'text': doc['text'][:2000],  # 减少到2000字符
                'filename': doc['filename'],
                'similarity': similarity
            })
    
    # 返回top N结果
    return sorted(scored_docs, key=lambda x: x['similarity'], reverse=True)[:max_results]

@app.get("/financial-analysis/stream")
async def financial_analysis_stream(
    analysis_type: str = "comprehensive",
    task_id: Optional[str] = None
):
    """流式财务分析接口"""
    try:
        print(f"\n=== 开始流式财务分析 ===")
        print(f"分析类型: {analysis_type}")
        print(f"任务ID: {task_id}")
        print(f"当前文档数量: {len(document_store)}")
        
        # 检查是否有上传的文档
        if len(document_store) == 0:
            async def error_stream():
                yield f"data: {json.dumps({'type': 'error', 'message': '请先上传财务数据文件'})}\n\n"
            
            return StreamingResponse(error_stream(), media_type="text/plain")
        
        # 收集所有文档的文本内容
        all_text = ""
        processed_files = []
        
        for doc_id, doc_info in document_store.items():
            all_text += f"\n\n=== 文件: {doc_info['filename']} ===\n"
            all_text += doc_info['text']
            processed_files.append(doc_info['filename'])
        
        print(f"处理文件: {processed_files}")
        print(f"总文本长度: {len(all_text)}")
        
        # 生成任务ID
        if not task_id:
            task_id = str(uuid.uuid4())
        
        # 导入财务分析服务并启动任务
        from financial_analysis import financial_service
        financial_service.start_task(task_id)
        
        # 返回流式响应
        return StreamingResponse(
            financial_service.analyze_financial_data_stream(all_text, analysis_type, task_id),
            media_type="text/plain",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
                "Access-Control-Allow-Headers": "Content-Type"
            }
        )
        
    except Exception as e:
        print(f"流式财务分析错误: {str(e)}")
        import traceback
        traceback.print_exc()
        
        async def error_stream(error_message: str):
            yield f"data: {json.dumps({'type': 'error', 'message': f'分析失败: {error_message}'})}\n\n"
        
        return StreamingResponse(error_stream(str(e)), media_type="text/plain")

@app.get("/documents")
async def get_documents():
    """获取所有已上传的文档内容"""
    try:
        if len(document_store) == 0:
            return {
                "success": False,
                "message": "没有已上传的文档"
            }
        
        documents = []
        for doc_id, doc_info in document_store.items():
            documents.append({
                "id": doc_id,
                "filename": doc_info['filename'],
                "text": doc_info['text'],
                "upload_time": doc_info.get('upload_time', ''),
                "file_size": len(doc_info['text'])
            })
        
        return {
            "success": True,
            "data": {
                "documents": documents,
                "total_count": len(documents)
            }
        }
        
    except Exception as e:
        print(f"获取文档错误: {str(e)}")
        return {
            "success": False,
            "message": f"获取文档失败: {str(e)}"
        }

if __name__ == "__main__":
    import uvicorn
    print("启动财务分析后端服务...")
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True) 